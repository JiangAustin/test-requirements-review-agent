from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .errors import REPORT_PARTIAL, ReviewError
from .models import (
    CheckResult,
    CheckStatus,
    Impact,
    ReportArtifacts,
    RequirementReview,
    ReviewReport,
    ScenarioResult,
    Severity,
    SourceRef,
)

_NO_SECTION_LABEL = "（无章节）"


def write_json(report: ReviewReport, path: Path) -> None:
    payload = json.dumps(
        report.model_dump(mode="json"),
        ensure_ascii=False,
        sort_keys=True,
        indent=2,
    ) + "\n"
    _write_text_atomic(path, payload)


def write_markdown(report: ReviewReport, path: Path) -> None:
    lines = [
        "# 需求评审报告",
        "",
        "## 执行摘要",
        f"- 共评审 {len(report.requirements)} 条需求。",
        f"- 阻塞问题 {len(_blocking_checks(report))} 项。",
        f"- 未完成项 {len(report.failures)} 项。",
        "",
        "## 指标定义",
        "- 需求可测试性得分（不是真实测试覆盖率）：只根据适用检查项是否已明确计算。",
        "- 建议场景覆盖度（不是真实测试覆盖率）：只表示建议测试场景对适用类别的覆盖程度。",
        "- 整体需求可测试性得分（不是真实测试覆盖率）："
        f"{_format_score(report.aggregate.testability)}",
        "- 整体建议场景覆盖度（不是真实测试覆盖率）："
        f"{_format_score(report.aggregate.scenario_coverage)}",
        "",
        "## 阻塞问题",
    ]
    lines.extend(_render_blocking_markdown(report))
    lines.extend(
        [
            "",
            "## 需求评审矩阵",
            "| 需求 ID | 需求内容 | 需求可测试性得分（不是真实测试覆盖率） "
            "| 建议场景覆盖度（不是真实测试覆盖率） | 手动测试缺口 | 自动化测试缺口 | 证据 |",
            "| --- | --- | --- | --- | --- | --- | --- |",
        ]
    )
    for requirement_review in report.requirements:
        manual_gap = _summarize_checks(requirement_review, Impact.MANUAL)
        automation_gap = _summarize_checks(requirement_review, Impact.AUTOMATION)
        evidence = _format_evidence_list(requirement_review.requirement.sources)
        lines.append(
            "| "
            f"{requirement_review.requirement.requirement_id} | "
            f"{_compact(requirement_review.requirement.text)} | "
            f"{_format_score(requirement_review.score.testability)} | "
            f"{_format_score(requirement_review.score.scenario_coverage)} | "
            f"{manual_gap} | {automation_gap} | {evidence} |"
        )

    lines.extend(["", "## 手动测试缺失信息"])
    lines.extend(_render_gap_markdown(report, Impact.MANUAL))
    lines.extend(["", "## 自动化测试缺失信息"])
    lines.extend(_render_gap_markdown(report, Impact.AUTOMATION))
    lines.extend(
        [
            "",
            "## 建议场景",
            "| 需求 ID | 场景类别 | 是否已覆盖 | 描述 | 证据 |",
            "| --- | --- | --- | --- | --- |",
        ]
    )
    for requirement_review in report.requirements:
        for scenario in requirement_review.analysis.scenarios:
            lines.append(
                "| "
                f"{requirement_review.requirement.requirement_id} | "
                f"{scenario.category} | "
                f"{_covered_label(scenario)} | "
                f"{_compact(scenario.description)} | "
                f"{_format_evidence_list(scenario.evidence)} |"
            )

    lines.extend(["", "## 未完成项"])
    lines.extend(_render_failures_markdown(report))
    lines.extend(
        [
            "",
            "## 运行元数据",
            "| 字段 | 值 |",
            "| --- | --- |",
            f"| schema_version | {report.schema_version} |",
            f"| run_id | {report.run_id or '无'} |",
            f"| generated_at | {report.generated_at} |",
            f"| provider_mode | {report.provider_mode.value} |",
            f"| model_name | {report.model_name or '无'} |",
            f"| rule_version | {report.rule_version} |",
        ]
    )
    _write_text_atomic(path, "\n".join(lines) + "\n")


def write_docx(report: ReviewReport, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    document = Document()
    document.core_properties.language = "zh-CN"
    document.add_heading("需求评审报告", level=0)

    document.add_heading("执行摘要", level=1)
    document.add_paragraph(f"共评审 {len(report.requirements)} 条需求。")
    document.add_paragraph(f"阻塞问题 {len(_blocking_checks(report))} 项。")
    document.add_paragraph(f"未完成项 {len(report.failures)} 项。")

    document.add_heading("指标定义", level=1)
    document.add_paragraph("需求可测试性得分（不是真实测试覆盖率）：只根据适用检查项是否已明确计算。")
    document.add_paragraph("建议场景覆盖度（不是真实测试覆盖率）：只表示建议测试场景对适用类别的覆盖程度。")
    document.add_paragraph(
        "整体需求可测试性得分（不是真实测试覆盖率）："
        f"{_format_score(report.aggregate.testability)}"
    )
    document.add_paragraph(
        "整体建议场景覆盖度（不是真实测试覆盖率）："
        f"{_format_score(report.aggregate.scenario_coverage)}"
    )

    document.add_heading("阻塞问题", level=1)
    for line in _render_blocking_lines(report):
        document.add_paragraph(line)

    document.add_heading("需求评审矩阵", level=1)
    matrix = document.add_table(rows=1, cols=7)
    matrix.style = "Table Grid"
    matrix.rows[0].cells[0].text = "需求 ID"
    matrix.rows[0].cells[1].text = "需求内容"
    matrix.rows[0].cells[2].text = "需求可测试性得分（不是真实测试覆盖率）"
    matrix.rows[0].cells[3].text = "建议场景覆盖度（不是真实测试覆盖率）"
    matrix.rows[0].cells[4].text = "手动测试缺口"
    matrix.rows[0].cells[5].text = "自动化测试缺口"
    matrix.rows[0].cells[6].text = "证据"
    for requirement_review in report.requirements:
        row = matrix.add_row().cells
        row[0].text = requirement_review.requirement.requirement_id
        row[1].text = requirement_review.requirement.text
        row[2].text = _format_score(requirement_review.score.testability)
        row[3].text = _format_score(requirement_review.score.scenario_coverage)
        row[4].text = _summarize_checks(requirement_review, Impact.MANUAL)
        row[5].text = _summarize_checks(requirement_review, Impact.AUTOMATION)
        row[6].text = _format_evidence_list(requirement_review.requirement.sources)

    document.add_heading("手动测试缺失信息", level=1)
    _add_gap_table(document, report, Impact.MANUAL)

    document.add_heading("自动化测试缺失信息", level=1)
    _add_gap_table(document, report, Impact.AUTOMATION)

    document.add_heading("建议场景", level=1)
    scenario_table = document.add_table(rows=1, cols=5)
    scenario_table.style = "Table Grid"
    scenario_table.rows[0].cells[0].text = "需求 ID"
    scenario_table.rows[0].cells[1].text = "场景类别"
    scenario_table.rows[0].cells[2].text = "是否已覆盖"
    scenario_table.rows[0].cells[3].text = "描述"
    scenario_table.rows[0].cells[4].text = "证据"
    for requirement_review in report.requirements:
        for scenario in requirement_review.analysis.scenarios:
            row = scenario_table.add_row().cells
            row[0].text = requirement_review.requirement.requirement_id
            row[1].text = scenario.category
            row[2].text = _covered_label(scenario)
            row[3].text = scenario.description
            row[4].text = _format_evidence_list(scenario.evidence)

    document.add_heading("未完成项", level=1)
    for line in _render_failure_lines(report):
        document.add_paragraph(line)

    document.add_heading("运行元数据", level=1)
    metadata_table = document.add_table(rows=1, cols=2)
    metadata_table.style = "Table Grid"
    metadata_table.rows[0].cells[0].text = "字段"
    metadata_table.rows[0].cells[1].text = "值"
    metadata_rows = (
        ("schema_version", report.schema_version),
        ("run_id", report.run_id or "无"),
        ("generated_at", report.generated_at),
        ("provider_mode", report.provider_mode.value),
        ("model_name", report.model_name or "无"),
        ("rule_version", report.rule_version),
    )
    for key, value in metadata_rows:
        row = metadata_table.add_row().cells
        row[0].text = key
        row[1].text = value

    try:
        document.save(str(temporary))
        with temporary.open("rb+") as stream:
            os.fsync(stream.fileno())
        temporary.replace(path)
    except Exception:
        if temporary.exists():
            temporary.unlink()
        raise


def render_all(report: ReviewReport, output_dir: Path) -> ReportArtifacts:
    output_dir.mkdir(parents=True, exist_ok=True)
    json_path = output_dir / "review.json"
    markdown_path = output_dir / "review.md"
    docx_path = output_dir / "review.docx"

    try:
        write_docx(report, docx_path)
    except (OSError, PackageNotFoundError) as error:
        updated_report = report.model_copy(
            update={
                "failures": report.failures
                + (
                    ReviewError(
                        code=REPORT_PARTIAL,
                        message="Word 报告生成失败，已保留 JSON 和 Markdown 产物。",
                        details={"error_type": type(error).__name__},
                    ),
                )
            }
        )
        write_json(updated_report, json_path)
        write_markdown(updated_report, markdown_path)
        return ReportArtifacts(json=json_path, markdown=markdown_path, docx=None, status="partial")

    write_json(report, json_path)
    write_markdown(report, markdown_path)
    return ReportArtifacts(
        json=json_path,
        markdown=markdown_path,
        docx=docx_path,
        status="complete",
    )


def _write_text_atomic(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    try:
        with temporary.open("w", encoding="utf-8", newline="\n") as stream:
            stream.write(content)
            stream.flush()
            os.fsync(stream.fileno())
        temporary.replace(path)
    except Exception:
        if temporary.exists():
            temporary.unlink()
        raise


def _render_blocking_markdown(report: ReviewReport) -> list[str]:
    lines = _render_blocking_lines(report)
    return lines if lines else ["- 无"]


def _render_blocking_lines(report: ReviewReport) -> list[str]:
    lines: list[str] = []
    for requirement_review, check in _blocking_checks(report):
        lines.append(
            "- "
            f"{requirement_review.requirement.requirement_id}: "
            f"{check.question or check.rationale or check.rule_id}；"
            f"证据：{_format_evidence_list(check.evidence)}"
        )
    return lines


def _blocking_checks(report: ReviewReport) -> list[tuple[RequirementReview, CheckResult]]:
    blocking: list[tuple[RequirementReview, CheckResult]] = []
    for requirement_review in report.requirements:
        for check in requirement_review.analysis.checks:
            if check.severity is Severity.BLOCKING and check.status in {
                CheckStatus.MISSING,
                CheckStatus.NEEDS_CONFIRMATION,
            }:
                blocking.append((requirement_review, check))
    return blocking


def _render_gap_markdown(report: ReviewReport, impact: Impact) -> list[str]:
    gap_lines = _render_gap_lines(report, impact)
    return gap_lines if gap_lines else ["- 无"]


def _render_gap_lines(report: ReviewReport, impact: Impact) -> list[str]:
    lines: list[str] = []
    for requirement_review in report.requirements:
        for check in _gap_checks(requirement_review, impact):
            lines.append(
                "- "
                f"{requirement_review.requirement.requirement_id}: "
                f"{check.question or check.rationale or check.rule_id}；"
                f"证据：{_format_evidence_list(check.evidence)}"
            )
    return lines


def _gap_checks(requirement_review: RequirementReview, impact: Impact) -> list[CheckResult]:
    matched: list[CheckResult] = []
    for check in requirement_review.analysis.checks:
        if check.status not in {CheckStatus.MISSING, CheckStatus.NEEDS_CONFIRMATION}:
            continue
        if impact is Impact.MANUAL and check.impact not in {Impact.MANUAL, Impact.BOTH}:
            continue
        if impact is Impact.AUTOMATION and check.impact not in {Impact.AUTOMATION, Impact.BOTH}:
            continue
        matched.append(check)
    return matched


def _render_failures_markdown(report: ReviewReport) -> list[str]:
    lines = _render_failure_lines(report)
    return lines if lines else ["- 无"]


def _render_failure_lines(report: ReviewReport) -> list[str]:
    lines: list[str] = []
    for failure in report.failures:
        details = json.dumps(failure.details, ensure_ascii=False, sort_keys=True)
        lines.append(
            f"- [{failure.code}] {failure.message}；详情：{details}"
        )
    return lines


def _format_score(value: float) -> str:
    return str(value)


def _covered_label(scenario: ScenarioResult) -> str:
    return "是" if scenario.covered else "否"


def _format_evidence(source: SourceRef) -> str:
    return f"p.{source.page} / {source.section or _NO_SECTION_LABEL}: “{source.quote}”"


def _format_evidence_list(evidence: tuple[SourceRef, ...]) -> str:
    if not evidence:
        return "无"
    return "；".join(_format_evidence(item) for item in evidence)


def _compact(value: str) -> str:
    return value.replace("\n", " ")


def _summarize_checks(requirement_review: RequirementReview, impact: Impact) -> str:
    checks = _gap_checks(requirement_review, impact)
    if not checks:
        return "无"
    return "；".join(check.question or check.rationale or check.rule_id for check in checks)


def _add_gap_table(document: DocxDocument, report: ReviewReport, impact: Impact) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "需求 ID"
    table.rows[0].cells[1].text = "规则 ID"
    table.rows[0].cells[2].text = "缺失信息"
    table.rows[0].cells[3].text = "证据"

    row_count = 1
    for requirement_review in report.requirements:
        for check in _gap_checks(requirement_review, impact):
            row = table.add_row().cells
            row[0].text = requirement_review.requirement.requirement_id
            row[1].text = check.rule_id
            row[2].text = check.question or check.rationale or check.rule_id
            row[3].text = _format_evidence_list(check.evidence)
            row_count += 1

    if row_count == 1:
        row = table.add_row().cells
        row[0].text = "无"
        row[1].text = "无"
        row[2].text = "无"
        row[3].text = "无"