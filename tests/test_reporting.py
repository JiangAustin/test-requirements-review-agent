from __future__ import annotations

import json
import re
from pathlib import Path

import pytest
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from requirements_review_agent.errors import REPORT_PARTIAL, ReviewError
from requirements_review_agent.models import (
    AggregateScore,
    AtomicRequirement,
    CheckResult,
    CheckStatus,
    FindingType,
    Impact,
    ProviderMode,
    RequirementAnalysis,
    RequirementReview,
    RequirementScore,
    ReviewReport,
    ScenarioResult,
    Severity,
    SourceRef,
)
from requirements_review_agent.reporting import render_all, write_json, write_markdown


def source(page: int, quote: str, *, section: str | None = None) -> SourceRef:
    return SourceRef(page=page, section=section, quote=quote)


def build_review_report() -> ReviewReport:
    wifi_source = source(
        3,
        "Device enters Wi-Fi pairing state within 30 seconds.",
        section="3.2 Connectivity",
    )
    timeout_source = source(5, "BLE reconnect timeout is 5 s.")

    requirement_one = RequirementReview(
        requirement=AtomicRequirement(
            requirement_id="REQ-CONN-001",
            text="设备应在 30 秒内进入 Wi-Fi pairing state。",
            sources=(wifi_source,),
            needs_manual_review=False,
        ),
        analysis=RequirementAnalysis(
            requirement_id="REQ-CONN-001",
            checks=(
                CheckResult(
                    rule_id="manual.precondition",
                    status=CheckStatus.MISSING,
                    impact=Impact.MANUAL,
                    severity=Severity.BLOCKING,
                    finding_type=FindingType.SUGGESTION,
                    evidence=(wifi_source,),
                    rationale="缺少进入 pairing state 前的设备前置条件。",
                    question="手动测试前设备需要处于什么初始状态？",
                    confidence=0.91,
                ),
                CheckResult(
                    rule_id="behavior.acceptance",
                    status=CheckStatus.COMPLETE,
                    impact=Impact.BOTH,
                    severity=Severity.NORMAL,
                    finding_type=FindingType.FACT,
                    evidence=(wifi_source,),
                    rationale="原文给出了进入 Wi-Fi pairing state 的时限。",
                    question=None,
                    confidence=0.97,
                ),
            ),
            scenarios=(
                ScenarioResult(
                    category="normal",
                    description="验证设备在 30 秒内进入 Wi-Fi pairing state。",
                    covered=True,
                    evidence=(wifi_source,),
                ),
                ScenarioResult(
                    category="recovery",
                    description="验证配网失败后的恢复路径。",
                    covered=False,
                    evidence=(),
                ),
            ),
        ),
        score=RequirementScore(
            requirement_id="REQ-CONN-001",
            testability=33.33,
            scenario_coverage=50.0,
        ),
    )

    requirement_two = RequirementReview(
        requirement=AtomicRequirement(
            requirement_id="REQ-BLE-002",
            text="BLE reconnect timeout 应可配置。",
            sources=(timeout_source,),
            needs_manual_review=True,
        ),
        analysis=RequirementAnalysis(
            requirement_id="REQ-BLE-002",
            checks=(
                CheckResult(
                    rule_id="automation.interface",
                    status=CheckStatus.MISSING,
                    impact=Impact.AUTOMATION,
                    severity=Severity.IMPORTANT,
                    finding_type=FindingType.SUGGESTION,
                    evidence=(timeout_source,),
                    rationale="未说明自动化可通过哪个接口配置 timeout。",
                    question="自动化测试通过哪个接口配置 BLE reconnect timeout？",
                    confidence=0.82,
                ),
            ),
            scenarios=(
                ScenarioResult(
                    category="boundary",
                    description="验证最小和最大 BLE reconnect timeout 配置。",
                    covered=True,
                    evidence=(timeout_source,),
                ),
            ),
        ),
        score=RequirementScore(
            requirement_id="REQ-BLE-002",
            testability=0.0,
            scenario_coverage=100.0,
        ),
    )

    return ReviewReport(
        schema_version="1.0",
        run_id="run-report-001",
        generated_at="2026-09-02T10:32:33Z",
        provider_mode=ProviderMode.COPILOT,
        model_name="gpt-5.4",
        rule_version="home-iot-v1",
        requirements=(requirement_one, requirement_two),
        aggregate=AggregateScore(testability=16.67, scenario_coverage=75.0),
        failures=(
            ReviewError(
                code="ANALYSIS_INVALID",
                message="存在 1 条需求需要人工确认。",
                details={"requirement_id": "REQ-BLE-002"},
            ),
        ),
    )


def docx_text(document: Document) -> str:
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    lines.append(cell.text)
    return "\n".join(lines)


def test_render_all_produces_complete_json_markdown_and_docx(tmp_path: Path) -> None:
    report = build_review_report()

    artifacts = render_all(report, tmp_path)

    assert artifacts.status == "complete"
    assert artifacts.json.exists()
    assert artifacts.markdown.exists()
    assert artifacts.docx is not None
    assert artifacts.docx.exists()

    json_text = artifacts.json.read_text(encoding="utf-8")
    markdown = artifacts.markdown.read_text(encoding="utf-8")
    document = Document(artifacts.docx)
    document_text = docx_text(document)

    assert document.core_properties.language == "zh-CN"
    assert document.paragraphs
    assert document.tables

    for requirement_id in ("REQ-CONN-001", "REQ-BLE-002"):
        assert requirement_id in json_text
        assert requirement_id in markdown
        assert requirement_id in document_text

    for score_text in ("33.33", "16.67", "100.0", "75.0"):
        assert score_text in json_text
        assert score_text in markdown
        assert score_text in document_text

    for shared_text in (
        "手动测试前设备需要处于什么初始状态？",
        "自动化测试通过哪个接口配置 BLE reconnect timeout？",
        "Device enters Wi-Fi pairing state within 30 seconds.",
        "BLE reconnect timeout is 5 s.",
        "Wi-Fi",
    ):
        assert shared_text in json_text
        assert shared_text in markdown
        assert shared_text in document_text

    for metric_label in (
        "建议场景覆盖度（不是真实测试覆盖率）",
        "需求可测试性得分（不是真实测试覆盖率）",
    ):
        assert metric_label in markdown
        assert metric_label in document_text


def test_markdown_contains_exact_nine_sections_and_stable_evidence_format(tmp_path: Path) -> None:
    report = build_review_report()
    path = tmp_path / "review.md"

    write_markdown(report, path)

    markdown = path.read_text(encoding="utf-8")
    assert re.findall(r"^##\s+(.+)$", markdown, flags=re.MULTILINE) == [
        "执行摘要",
        "指标定义",
        "阻塞问题",
        "需求评审矩阵",
        "手动测试缺失信息",
        "自动化测试缺失信息",
        "建议场景",
        "未完成项",
        "运行元数据",
    ]
    assert (
        'p.3 / 3.2 Connectivity: “Device enters Wi-Fi pairing state within 30 seconds.”'
        in markdown
    )
    assert 'p.5 / （无章节）: “BLE reconnect timeout is 5 s.”' in markdown


def test_markdown_escapes_dynamic_table_cells(tmp_path: Path) -> None:
    report = build_review_report()
    original = report.requirements[0]
    special_source = source(7, "Mode A | Mode B\nremains available.", section="7.1")
    updated_requirement = original.requirement.model_copy(
        update={
            "text": "支持 Mode A | Mode B。\n切换后状态保持。",
            "sources": (special_source,),
        }
    )
    updated_review = original.model_copy(update={"requirement": updated_requirement})
    updated_report = report.model_copy(
        update={"requirements": (updated_review, *report.requirements[1:])}
    )
    path = tmp_path / "review.md"

    write_markdown(updated_report, path)

    markdown = path.read_text(encoding="utf-8")
    matrix_line = next(
        line for line in markdown.splitlines() if line.startswith("| REQ-CONN-001 |")
    )
    assert "支持 Mode A \\| Mode B。<br>切换后状态保持。" in matrix_line
    assert "Mode A \\| Mode B<br>remains available." in matrix_line


def test_write_json_matches_model_dump_text_exactly(tmp_path: Path) -> None:
    report = build_review_report()
    path = tmp_path / "review.json"

    write_json(report, path)

    expected = json.dumps(
        report.model_dump(mode="json"),
        ensure_ascii=False,
        sort_keys=True,
        indent=2,
    ) + "\n"
    assert path.read_text(encoding="utf-8") == expected


@pytest.mark.parametrize("error", [OSError("disk full"), PackageNotFoundError("pkg")])
def test_docx_failure_returns_partial_artifacts_and_report_partial_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    error: Exception,
) -> None:
    def fail_docx(report: ReviewReport, path: Path) -> None:
        raise error

    monkeypatch.setattr("requirements_review_agent.reporting.write_docx", fail_docx)

    artifacts = render_all(build_review_report(), tmp_path)

    assert artifacts.status == "partial"
    assert artifacts.docx is None
    assert artifacts.json.exists()
    assert artifacts.markdown.exists()

    json_text = artifacts.json.read_text(encoding="utf-8")
    markdown = artifacts.markdown.read_text(encoding="utf-8")
    for text in (REPORT_PARTIAL, '"error_type":', error.__class__.__name__):
        assert text in json_text
        assert text in markdown
    assert "disk full" not in json_text
    assert "disk full" not in markdown
    assert str(tmp_path) not in json_text
    assert str(tmp_path) not in markdown


def test_render_all_does_not_catch_unexpected_docx_error(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def fail_docx(report: ReviewReport, path: Path) -> None:
        raise RuntimeError("unexpected")

    monkeypatch.setattr("requirements_review_agent.reporting.write_docx", fail_docx)

    with pytest.raises(RuntimeError, match="unexpected"):
        render_all(build_review_report(), tmp_path)


@pytest.mark.parametrize("failed_writer", ["write_json", "write_markdown"])
def test_text_report_failure_does_not_leave_orphan_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    failed_writer: str,
) -> None:
    complete = render_all(build_review_report(), tmp_path)
    assert complete.docx is not None and complete.docx.exists()

    def fail_writer(report: ReviewReport, path: Path) -> None:
        raise OSError("disk full")

    monkeypatch.setattr(f"requirements_review_agent.reporting.{failed_writer}", fail_writer)

    with pytest.raises(OSError, match="disk full"):
        render_all(build_review_report(), tmp_path)

    assert not (tmp_path / "review.docx").exists()


def test_partial_rerender_removes_stale_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    complete = render_all(build_review_report(), tmp_path)
    assert complete.docx is not None and complete.docx.exists()

    def fail_docx(report: ReviewReport, path: Path) -> None:
        raise OSError("disk full")

    monkeypatch.setattr("requirements_review_agent.reporting.write_docx", fail_docx)

    partial = render_all(build_review_report(), tmp_path)

    assert partial.status == "partial"
    assert partial.docx is None
    assert not (tmp_path / "review.docx").exists()


def test_markdown_write_preserves_existing_target_when_replace_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "review.md"
    old_report = build_review_report()
    new_report = old_report.model_copy(update={"generated_at": "2026-09-02T11:00:00Z"})

    write_markdown(old_report, path)
    original = path.read_text(encoding="utf-8")
    original_replace = Path.replace

    def fail_replace(self: Path, target: Path) -> Path:
        if self.name == "review.md.tmp":
            raise OSError("disk full")
        return original_replace(self, target)

    monkeypatch.setattr(Path, "replace", fail_replace)

    with pytest.raises(OSError, match="disk full"):
        write_markdown(new_report, path)

    assert path.read_text(encoding="utf-8") == original
    assert not (tmp_path / "review.md.tmp").exists()