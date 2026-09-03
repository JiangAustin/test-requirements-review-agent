from __future__ import annotations

import json
from pathlib import Path

import fitz
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from requirements_review_agent.analysis import AnalysisBatch
from requirements_review_agent.models import (
    AnalysisSubmission,
    CheckResult,
    CheckStatus,
    FindingType,
    Impact,
    ProviderMode,
    RequirementAnalysis,
    ReviewReport,
    Severity,
    SourceRef,
)
from requirements_review_agent.service import ReviewService


def build_text_table_pdf(path: Path) -> Path:
    document = fitz.open()
    page = document.new_page()
    page.insert_textbox(
        fitz.Rect(50, 50, 430, 220),
        "Connectivity:\n"
        "1. Device enters Wi-Fi pairing state within 30 seconds.\n"
        "2. BLE reconnect timeout is 5 s and should recover automatically.",
        fontsize=12,
    )

    x0, y0, width, height = 50, 170, 320, 90
    col_width = width / 2
    row_height = height / 2
    for row_index in range(3):
        page.draw_line(
            (x0, y0 + row_index * row_height),
            (x0 + width, y0 + row_index * row_height),
        )
    for column_index in range(3):
        page.draw_line(
            (x0 + column_index * col_width, y0),
            (x0 + column_index * col_width, y0 + height),
        )

    page.insert_text((x0 + 10, y0 + 18), "Param")
    page.insert_text((x0 + col_width + 10, y0 + 18), "Value")
    page.insert_text((x0 + 10, y0 + row_height + 18), "Timeout")
    page.insert_text((x0 + col_width + 10, y0 + row_height + 18), "30 s")
    document.save(path)
    document.close()
    return path


def parse_frontmatter(path: Path) -> tuple[dict[str, object], str]:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    assert len(lines) >= 3
    assert lines[0] == "---"
    end_index = lines.index("---", 1)
    metadata: dict[str, object] = {}
    for line in lines[1:end_index]:
        key, raw_value = line.split(":", 1)
        value = raw_value.strip()
        if value.startswith("[") and value.endswith("]"):
            items = [item.strip().strip("'\"") for item in value[1:-1].split(",") if item.strip()]
            metadata[key.strip()] = items
        else:
            metadata[key.strip()] = value.strip("'\"")
    body = "\n".join(lines[end_index + 1 :]).strip()
    return metadata, body


def requirement_source(requirement: object) -> SourceRef:
    return SourceRef.model_validate(requirement.sources[0].model_dump(mode="json"))


def build_submission(batch: AnalysisBatch) -> AnalysisSubmission:
    analyses: list[RequirementAnalysis] = []
    for requirement in batch.requirements:
        applicable_rules = batch.applicable[requirement.requirement_id]
        evidence = (requirement_source(requirement),)
        scenarios: list[dict[str, object]] = []
        categories: list[str] = []
        checks: list[CheckResult] = []
        for index, rule in enumerate(applicable_rules):
            if rule.scenario_category and rule.scenario_category not in categories:
                categories.append(rule.scenario_category)
                scenarios.append(
                    {
                        "category": rule.scenario_category,
                        "description": f"验证 {rule.scenario_category} 场景。",
                        "covered": True,
                        "evidence": [item.model_dump(mode="json") for item in evidence],
                    }
                )

            if index == 0:
                status = CheckStatus.COMPLETE
                severity = Severity.NORMAL
                question = None
                rationale = "原文已明确说明。"
            elif index % 2 == 0:
                status = CheckStatus.NEEDS_CONFIRMATION
                severity = Severity.IMPORTANT
                question = f"请确认 {rule.rule_id} 相关信息是否完整。"
                rationale = "原文存在需人工确认的信息。"
            else:
                status = CheckStatus.MISSING
                severity = Severity.BLOCKING
                question = f"请补充 {rule.rule_id} 相关缺失信息。"
                rationale = "原文缺少必要细节。"

            checks.append(
                CheckResult(
                    rule_id=rule.rule_id,
                    status=status,
                    impact=rule.impact,
                    severity=severity,
                    finding_type=FindingType.FACT,
                    evidence=evidence,
                    rationale=rationale,
                    question=question,
                    confidence=0.93,
                )
            )

        analyses.append(
            RequirementAnalysis.model_validate(
                {
                    "requirement_id": requirement.requirement_id,
                    "checks": [check.model_dump(mode="json") for check in checks],
                    "scenarios": scenarios,
                }
            )
        )

    return AnalysisSubmission(schema_version="1.0", requirements=tuple(analyses))


def collect_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def markdown_table_rows(markdown: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in markdown.splitlines():
        if not line.startswith("|"):
            continue
        cells: list[str] = []
        current: list[str] = []
        escaped = False
        for char in line[1:]:
            if escaped:
                current.append(char)
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == "|":
                cells.append("".join(current).strip().replace("<br>", "\n"))
                current = []
            else:
                current.append(char)
        if cells and not all(set(cell) <= {"-", ":"} for cell in cells):
            rows.append(cells)
    return rows


def docx_tables_by_heading(path: Path) -> dict[str, list[list[str]]]:
    document = Document(path)
    tables: dict[str, list[list[str]]] = {}
    heading: str | None = None
    for block in document.iter_inner_content():
        if isinstance(block, Paragraph) and block.style.name.startswith("Heading"):
            heading = block.text
        elif isinstance(block, Table) and heading is not None:
            tables[heading] = [[cell.text for cell in row.cells] for row in block.rows]
    return tables


def markdown_sections(markdown: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current: list[str] | None = None
    for line in markdown.splitlines():
        if line.startswith("## "):
            current = sections.setdefault(line[3:], [])
        elif current is not None:
            current.append(line)
    return sections


def evidence_text(evidence: tuple[SourceRef, ...]) -> str:
    return "；".join(
        f"p.{source.page} / {source.section or '（无章节）'}: “{source.quote}”"
        for source in evidence
    ) or "无"


def test_workspace_agent_and_portable_config_contracts_exist() -> None:
    root = Path.cwd()
    agent_path = root / ".github" / "agents" / "requirements-review.agent.md"
    mcp_path = root / ".vscode" / "mcp.json"
    env_example_path = root / ".env.example"
    readme_path = root / "README.md"

    assert agent_path.exists()
    assert mcp_path.exists()
    assert env_example_path.exists()
    assert readme_path.exists()

    frontmatter, body = parse_frontmatter(agent_path)
    assert frontmatter == {
        "name": "Requirements Review",
        "description": "Review PDF requirements for manual and automation test gaps.",
        "tools": ["read", "requirements-review/*"],
    }

    ordered_steps = [
        "1. 从 Chat attachment metadata 读取 `requirements/` 下的 PDF 路径；"
        "没有附件时询问该目录内的 PDF 路径",
        "2. 说明 provider 和数据去向，并在非 Copilot 外部传输前取得明确确认",
        "3. 调用 prepare_review，并在 unsupported/scanned/encrypted PDF 时停止",
        "4. Copilot 模式下对每个 batch 调用 get_analysis_batch、"
        "生成 schema-valid 结果并调用 submit_analysis",
        "5. company_api 或 local 模式调用 run_provider_analysis",
        "6. 调用 finalize_review，然后调用 get_review_status",
        "7. 用中文总结 blocking findings、两个指标定义、failed items 和本地产物路径",
    ]
    last_index = -1
    for step in ordered_steps:
        current_index = body.index(step)
        assert current_index > last_index
        last_index = current_index

    assert "rule pack 默认 `home-iot-v1`" in body
    assert "model mode 默认 `copilot`" in body
    assert "需求可测试性得分不是现有用例覆盖率，也不是测试执行覆盖率" in body
    assert "建议场景覆盖度不是现有用例覆盖率，也不是测试执行覆盖率" in body
    assert "company_api" in body and "显式确认 external transfer" in body
    assert "local" in body and "数据不离开本机" in body
    assert "review PDF" in body
    assert "manual" in body
    assert "automation" in body

    mcp_config = json.loads(mcp_path.read_text(encoding="utf-8"))
    assert mcp_config == {
        "servers": {
            "requirements-review": {
                "type": "stdio",
                "command": "uv",
                "args": ["run", "requirements-review-mcp"],
            }
        }
    }
    encoded_mcp = json.dumps(mcp_config, ensure_ascii=False)
    assert "C:\\\\" not in encoded_mcp
    assert "API_KEY" not in encoded_mcp
    assert "token" not in encoded_mcp.lower()

    assert env_example_path.read_text(encoding="utf-8").splitlines() == [
        "RRA_COMPANY_BASE_URL=",
        "RRA_COMPANY_API_KEY=",
        "RRA_COMPANY_MODEL=",
        "RRA_LOCAL_BASE_URL=",
        "RRA_LOCAL_MODEL=",
    ]


def test_text_table_pdf_produces_consistent_local_artifacts(tmp_path: Path) -> None:
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    rules_dir = workspace / "rules"
    rules_dir.mkdir()
    rule_pack = Path("rules/home-iot-v1.yaml").read_text(encoding="utf-8")
    rule_pack += """

  - id: manual.operator_observation
    question: "请确认 manual operator observation 的步骤与预期结果。"
    weight: 3
    impact: manual
    scenario_category: manual-observation
    always: true
    keywords: []
"""
    rules_dir.joinpath("home-iot-v1.yaml").write_text(
        rule_pack,
        encoding="utf-8",
    )
    pdf_path = build_text_table_pdf(workspace / "input.pdf")
    service = ReviewService(workspace)

    prepared = service.prepare(pdf_path, "home-iot-v1", ProviderMode.COPILOT)
    batch = service.get_batch(prepared.run_id, 0)
    submission = build_submission(batch)
    analyzed = service.submit(prepared.run_id, submission)
    artifacts = service.finalize(prepared.run_id)
    status = service.status(prepared.run_id)

    assert analyzed.stage == "analyzed"
    assert status.stage == "finalized"
    assert status.artifacts is not None
    assert status.artifacts.json.exists()
    assert status.artifacts.markdown.exists()
    assert status.artifacts.docx is not None
    assert status.artifacts.docx.exists()

    report = ReviewReport.model_validate_json(artifacts.json.read_text(encoding="utf-8"))
    markdown_text = artifacts.markdown.read_text(encoding="utf-8")
    docx_text = collect_docx_text(status.artifacts.docx)
    markdown_rows = markdown_table_rows(markdown_text)
    sections = markdown_sections(markdown_text)
    docx_tables = docx_tables_by_heading(status.artifacts.docx)

    assert set(report.aggregate.model_dump(mode="json")) == {"testability", "scenario_coverage"}
    assert report.aggregate.testability >= 0
    assert report.aggregate.scenario_coverage >= 0
    assert any(item.sources[0].table_index is None for item in batch.requirements)
    assert any(item.sources[0].table_index is not None for item in batch.requirements)
    assert any(review.requirement.sources[0].table_index is None for review in report.requirements)
    assert any(
        review.requirement.sources[0].table_index is not None
        for review in report.requirements
    )
    assert all(
        check.evidence
        for review in report.requirements
        for check in review.analysis.checks
        if check.finding_type is FindingType.FACT
    )
    assert all(
        check.question and any("\u4e00" <= char <= "\u9fff" for char in check.question)
        for review in report.requirements
        for check in review.analysis.checks
        if check.status in {CheckStatus.MISSING, CheckStatus.NEEDS_CONFIRMATION}
    )
    assert {Impact.MANUAL, Impact.AUTOMATION, Impact.BOTH} <= {
        check.impact
        for review in report.requirements
        for check in review.analysis.checks
        if check.status in {CheckStatus.MISSING, CheckStatus.NEEDS_CONFIRMATION}
    }
    assert any(
        source.table_index is not None
        for review in report.requirements
        for source in review.requirement.sources
    )

    for review in report.requirements:
        assert review.requirement.requirement_id in markdown_text
        assert review.requirement.requirement_id in docx_text
        assert str(review.score.testability) in markdown_text
        assert str(review.score.testability) in docx_text
        assert str(review.score.scenario_coverage) in markdown_text
        assert str(review.score.scenario_coverage) in docx_text
        requirement_evidence = evidence_text(review.requirement.sources)
        expected_matrix_row = [
            review.requirement.requirement_id,
            review.requirement.text,
            str(review.score.testability),
            str(review.score.scenario_coverage),
        ]
        assert any(
            row[:4] == expected_matrix_row and row[6] == requirement_evidence
            for row in markdown_rows
            if len(row) == 7
        )
        assert any(
            row[:4] == expected_matrix_row and row[6] == requirement_evidence
            for row in docx_tables["需求评审矩阵"]
            if len(row) == 7
        )
        for check in review.analysis.checks:
            if check.status in {CheckStatus.MISSING, CheckStatus.NEEDS_CONFIRMATION}:
                assert check.question is not None
                expected_gap_row = [
                    review.requirement.requirement_id,
                    check.rule_id,
                    check.question,
                    evidence_text(check.evidence),
                ]
                expected_gap_line = (
                    f"- {review.requirement.requirement_id}: {check.question}；"
                    f"证据：{evidence_text(check.evidence)}"
                )
                if check.impact in {Impact.MANUAL, Impact.BOTH}:
                    assert expected_gap_line in sections["手动测试缺失信息"]
                if check.impact in {Impact.AUTOMATION, Impact.BOTH}:
                    assert expected_gap_line in sections["自动化测试缺失信息"]
                if check.impact in {Impact.MANUAL, Impact.BOTH}:
                    assert expected_gap_row in docx_tables["手动测试缺失信息"]
                if check.impact in {Impact.AUTOMATION, Impact.BOTH}:
                    assert expected_gap_row in docx_tables["自动化测试缺失信息"]
        for scenario in review.analysis.scenarios:
            expected_scenario_row = [
                review.requirement.requirement_id,
                scenario.category,
                "是" if scenario.covered else "否",
                scenario.description,
                evidence_text(scenario.evidence),
            ]
            assert expected_scenario_row in markdown_rows
            assert expected_scenario_row in docx_tables["建议场景"]

    assert "需求可测试性得分（不是真实测试覆盖率）" in markdown_text
    assert "建议场景覆盖度（不是真实测试覆盖率）" in markdown_text
    assert "不表示现有用例覆盖率，也不表示测试执行覆盖率" in markdown_text
    assert "需求可测试性得分（不是真实测试覆盖率）" in docx_text
    assert "建议场景覆盖度（不是真实测试覆盖率）" in docx_text
    assert "不表示现有用例覆盖率，也不表示测试执行覆盖率" in docx_text
    assert status.artifacts.json.is_relative_to(workspace)
    assert status.artifacts.markdown.is_relative_to(workspace)
    assert status.artifacts.docx.is_relative_to(workspace)


def test_readme_defines_setup_security_and_metric_boundaries() -> None:
    readme = Path("README.md").read_text(encoding="utf-8")

    required_fragments = [
        "Python >=3.12,<3.14",
        "uv sync --dev",
        "uv run pytest",
        "uv run ruff check .",
        "uv run mypy src",
        "uv run mcp dev src/requirements_review_agent/server.py",
        "repository root",
        "Requirements Review",
        "MCP",
        "Copilot",
        "prepare_review",
        "get_analysis_batch",
        "submit_analysis",
        "run_provider_analysis",
        "finalize_review",
        "get_review_status",
        "RRA_COMPANY_BASE_URL",
        "RRA_COMPANY_API_KEY",
        "RRA_COMPANY_MODEL",
        "RRA_LOCAL_BASE_URL",
        "RRA_LOCAL_MODEL",
        ".runs/",
        ".env",
        "input PDFs",
        "JSON",
        "Markdown",
        "DOCX",
        "PDF_ENCRYPTED",
        "PDF_DAMAGED",
        "PDF_SCANNED",
        "PDF_OUTSIDE_WORKSPACE",
        "RULE_PACK_INVALID",
        "ANALYSIS_INVALID",
        "PROVIDER_UNAVAILABLE",
        "REPORT_PARTIAL",
        "no OCR",
        "Problems panel",
        "synthetic",
        "建议场景覆盖度不是现有用例覆盖率，也不是测试执行覆盖率",
        "需求可测试性得分不是现有用例覆盖率，也不是测试执行覆盖率",
        "显式确认",
        "secrets never",
        "不会自动加载 .env",
        '$env:RRA_COMPANY_API_KEY = Read-Host',
        "local 模式的数据不离开本机",
    ]
    for fragment in required_fragments:
        assert fragment in readme